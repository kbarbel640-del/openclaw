from docx import Document
import glob

files = [
    "/Users/sulaxd/Documents/ITRI_AI_Course/課程DM-無程式 AI 社群小編養成術-自動化內容生產全攻略(115年6月 )(v20260112).doc",
    "/Users/sulaxd/Documents/ITRI_AI_Course/課程DM-無程式AI知識管家實務-文件整理、LINE智能問答全攻略(115年6月)(v20260112).doc",
    "/Users/sulaxd/Documents/ITRI_AI_Course/課程DM-零程式 AI 編輯術-情資蒐集、自動分類與內容發布全攻略(115年5月)(v20260112).doc",
]

# Try to find files with glob if exact names don't work
import os
itri_dir = "/Users/sulaxd/Documents/ITRI_AI_Course/"
print("=== Files in ITRI dir ===")
for f in os.listdir(itri_dir):
    print(f)

print("\n=== Reading .doc files ===")
for f in os.listdir(itri_dir):
    if f.endswith('.doc') or f.endswith('.docx'):
        filepath = os.path.join(itri_dir, f)
        print(f"\n\n{'='*60}")
        print(f"FILE: {f}")
        print('='*60)
        try:
            doc = Document(filepath)
            for p in doc.paragraphs:
                if p.text.strip():
                    print(p.text)
            for t in doc.tables:
                for row in t.rows:
                    cells = [cell.text.replace("\n"," ")[:80] for cell in row.cells]
                    print(" | ".join(cells))
                print("---")
        except Exception as e:
            print(f"Error: {e}")
