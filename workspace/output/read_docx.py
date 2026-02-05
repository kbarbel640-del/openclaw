from docx import Document
doc = Document("/Users/sulaxd/Downloads/115年上學期新編科目課程大綱_20250105v1 (1).docx")
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text)
print("\n=== TABLES ===")
for t in doc.tables:
    for row in t.rows:
        cells = [cell.text.replace("\n"," ") for cell in row.cells]
        print(" | ".join(cells))
    print("---")
